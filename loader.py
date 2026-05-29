"""从 data 目录加载文档。"""

from pathlib import Path
from typing import TypedDict

from pypdf import PdfReader

import config
from logger import get_logger

log = get_logger()


class Document(TypedDict):
    text: str
    source: str
    doc_id: str


from multimodal import IMAGE_EXTENSIONS, describe_image_file, extract_pdf_image_descriptions, merge_text_and_vision
from video_loader import VIDEO_EXTENSIONS, read_video_as_text

SUPPORTED_EXTENSIONS = (
    {".txt", ".md", ".pdf", ".docx", ".html", ".htm"} | IMAGE_EXTENSIONS | VIDEO_EXTENSIONS
)


def _read_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def _read_pdf(path: Path) -> str:
    reader = PdfReader(str(path))
    parts: list[str] = []
    for i, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        if text.strip():
            parts.append(f"[第 {i} 页]\n{text}")
    extracted = "\n\n".join(parts)
    if config.OCR_ENABLED:
        from pdf_ocr import pdf_text_with_ocr_fallback

        extracted = pdf_text_with_ocr_fallback(path, extracted, config.OCR_MIN_TEXT_LEN)

    if config.USE_MULTIMODAL:
        vision = extract_pdf_image_descriptions(path)
        extracted = merge_text_and_vision(extracted, vision)
    return extracted


def _read_docx(path: Path) -> str:
    try:
        from docx import Document as DocxDocument
    except ImportError as e:
        raise ImportError("读取 docx 需要安装: pip install python-docx") from e
    doc = DocxDocument(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


def _read_html(path: Path) -> str:
    try:
        from bs4 import BeautifulSoup
    except ImportError as e:
        raise ImportError("读取 html 需要安装: pip install beautifulsoup4") from e
    html = path.read_text(encoding="utf-8")
    soup = BeautifulSoup(html, "html.parser")
    for tag in soup(["script", "style", "nav", "footer"]):
        tag.decompose()
    return soup.get_text(separator="\n", strip=True)


def _read_image(path: Path) -> str:
    if not config.USE_MULTIMODAL:
        log.warning("跳过图片 %s（USE_MULTIMODAL=false）", path.name)
        return ""
    return describe_image_file(path)


def _read_file(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix in IMAGE_EXTENSIONS:
        return _read_image(path)
    if suffix in VIDEO_EXTENSIONS:
        return read_video_as_text(path)
    if suffix == ".pdf":
        return _read_pdf(path)
    if suffix == ".docx":
        return _read_docx(path)
    if suffix in (".html", ".htm"):
        return _read_html(path)
    return _read_text_file(path)


def load_documents(
    data_dir: Path,
    sources: set[str] | None = None,
) -> list[Document]:
    """遍历 data_dir，加载文档；sources 为相对路径集合时仅加载这些文件。"""
    documents: list[Document] = []
    if not data_dir.exists():
        return documents

    files = sorted(
        p for p in data_dir.rglob("*") if p.is_file() and p.suffix.lower() in SUPPORTED_EXTENSIONS
    )

    for path in files:
        rel_source = str(path.relative_to(data_dir)).replace("\\", "/")
        if sources is not None and rel_source not in sources:
            continue

        try:
            text = _read_file(path)
        except Exception as e:
            log.warning("跳过 %s: %s", path.name, e)
            continue

        text = text.strip()
        if not text:
            log.warning("跳过空文件: %s", path.name)
            continue

        documents.append(
            Document(
                text=text,
                source=rel_source,
                doc_id=rel_source,
            )
        )

    return documents
